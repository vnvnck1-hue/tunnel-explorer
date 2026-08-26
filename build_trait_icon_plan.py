from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE

OUT = r"C:\Users\Loadcomplete\Documents\ChatGPT\땅굴 크루 만들기\docs\tunnel-crew-trait-icon-production-plan-v1.docx"

NAVY = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
GOLD = "A56B00"
INK = "17212B"
MUTED = "5B6770"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tcPr.append(shd)
    shd.set(qn('w:fill'), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in('w:tcMar')
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar')
        tcPr.append(tcMar)
    for m, v in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tcMar.find(qn(f'w:{m}'))
        if node is None:
            node = OxmlElement(f'w:{m}')
            tcMar.append(node)
        node.set(qn('w:w'), str(v))
        node.set(qn('w:type'), 'dxa')


def set_cell_width(cell, width_dxa):
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = tcPr.find(qn('w:tcW'))
    if tcW is None:
        tcW = OxmlElement('w:tcW')
        tcPr.append(tcW)
    tcW.set(qn('w:w'), str(width_dxa))
    tcW.set(qn('w:type'), 'dxa')


def set_table_geometry(table, widths):
    table.autofit = False
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = OxmlElement('w:tblW')
        tblPr.append(tblW)
    tblW.set(qn('w:w'), str(sum(widths)))
    tblW.set(qn('w:type'), 'dxa')
    tblInd = tblPr.find(qn('w:tblInd'))
    if tblInd is None:
        tblInd = OxmlElement('w:tblInd')
        tblPr.append(tblInd)
    tblInd.set(qn('w:w'), '120')
    tblInd.set(qn('w:type'), 'dxa')
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for w in widths:
        col = OxmlElement('w:gridCol')
        col.set(qn('w:w'), str(w))
        grid.append(col)
    for row in table.rows:
        for cell, w in zip(row.cells, widths):
            set_cell_width(cell, w)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), 'true')
    trPr.append(tblHeader)


def set_keep_with_next(paragraph, value=True):
    pPr = paragraph._p.get_or_add_pPr()
    node = pPr.find(qn('w:keepNext'))
    if node is None:
        node = OxmlElement('w:keepNext')
        pPr.append(node)
    node.set(qn('w:val'), 'true' if value else 'false')


def add_text(p, text, bold=False, color=None, size=None, italic=False):
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    if size:
        r.font.size = Pt(size)
    r.font.name = 'Calibri'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    return r


def add_para(doc, text='', style=None, before=0, after=6, line=1.25, color=INK, size=11, bold=False):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    if text:
        add_text(p, text, bold=bold, color=color, size=size)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f'Heading {level}')
    p.paragraph_format.keep_with_next = True
    add_text(p, text, bold=True, color=BLUE if level < 3 else DARK_BLUE, size={1:16,2:13,3:12}[level])
    return p


def add_callout(doc, label, body):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    add_text(p, label + '  ', bold=True, color=GOLD, size=10)
    add_text(p, body, color=INK, size=10)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def make_table(doc, headers, rows, widths, font_size=9):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, widths)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for cell, text in zip(hdr.cells, headers):
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        add_text(p, text, bold=True, color=NAVY, size=font_size)
    for ridx, row in enumerate(rows):
        cells = table.add_row().cells
        for cidx, (cell, text) in enumerate(zip(cells, row)):
            if ridx % 2 == 1:
                set_cell_shading(cell, 'FAFBFC')
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            add_text(p, str(text), color=INK, size=font_size)
    return table


base_groups = [
    ('01', '드릴 모터·비트', '기계식 드릴 / 회전축', '토크 증폭 모터, 단층 추적 비트, 과압 실린더, 다이아몬드 비트, 굴진 과급', '이미 제작됨: 드릴러 토크 증폭 모터. 속도·압력·피해량은 작은 링/게이지 오버레이로 구분'),
    ('02', '균열·단층', '균열 유지 / 영구 균열', '균열 고착 수지, 열회수 균열기, 영구 단층 각인, 단층 파쇄축', '수지·열회수·영구 효과는 색상 또는 작은 잠금/열 아이콘으로 구분'),
    ('03', '코어 암반', '맨틀 / 코어 암반', '맨틀 천공 키', '고유 의미가 분명해 별도 아이콘 권장'),
    ('04', '지진 충격파', '지층 붕괴 / 광역 파쇄', '지진 공진축, 지각 붕괴 코어', '동심 충격파 링의 수로 범위·강도 표현'),
    ('05', '회전 약실·급탄', '재사용 대기시간 / 탄창 / 급탄', '고속 약실 순환기, 무정지 발사 사이클, 다중 급탄, 병렬 급탄기', '이미 제작됨: 거너 고속 약실 순환기. 연사·탄창 효과는 속도선/탄환 수 오버레이'),
    ('06', '파쇄탄·폭파 작약', '벽 피해 / 반경 / 군집 폭발', '텅스텐 라이너, 파쇄 확장 슬리브, 군집 파쇄 작약, 파쇄 탄막, 군집 붕괴탄', '폭발 링, 벽 조각, 탄두 형태를 작은 오버레이로 변형'),
    ('07', '신관·기폭기', '지연 / 원격 기폭', '가변 지연 신관, 원격 과충전 기폭기', '시계·무선 신호 마크만 추가하면 됨'),
    ('08', '파편·산탄', '대인 피해 / 파편탄', '대인 파편 재킷, 심층 지배망', '심층 지배망은 파편 + 경험치/벽 피해/탄창 보조 마크 조합'),
    ('09', '플레어·광원', '조명 반경 / 지속시간', '고광도 연소제, 원격 측량 릴레이, 인공 태양 플레어', '이미 제작됨: 스카우트 고광도 연소제. 광원 크기·태양 링으로 상위 효과 표현'),
    ('10', '그래플·윈치', '거리 / 재사용 대기시간 / 이동', '경량 인장 케이블, 회생 제동 윈치, 수평선 견인 장치', '케이블 길이, 윈치 링, 수평선 화살표 오버레이'),
    ('11', '스캐너·레이더', '정찰 반경 / 시야 / 탐사', '다중 스펙트럼 스캐너, 감시자 광학계, 광맥 탐지망', '이미지 본체는 스캐너로 공유; 레이더 링·광맥 점·경험치 보조 마크 사용'),
    ('12', '축전지·전력 노드', '지속시간 / 공급 반경 / 자율 전력', '고밀도 축전지, 분산 전력 프로토콜, 폐쇄형 자율 동력로, 심층 요새 네트워크', '이미 제작됨: 엔지니어 고밀도 축전지. 노드 수·반경은 연결선/노드 점으로 표현'),
    ('13', '센트리 터렛', '탄창 / 냉각 / 설치 수 / 중량화', '확장 탄약 호퍼, 능동 냉각 재킷, 쌍중 센트리 설계, 중량 레일 센트리', '터렛 본체 하나로 공유하고 탄창·냉각·쌍열·중량 부품을 오버레이'),
    ('14', '쌍동력·동시 운용', '드릴·총 동시 강화', '쌍동력 융합, 쌍축 액추에이터, 무제한 출력', '두 개의 회전축 또는 교차 화살표로 통일'),
    ('15', '도탄', '도탄 횟수', '무한 도탄', '단독 사용. 화살표 궤적 2~3개로 강화 단계 표현'),
    ('16', '생존 격벽', '최대 체력 / 회복', '생존 격벽', '단독 사용. 방어판 + 하트 보조 마크'),
    ('17', '관통 굴착탄', '관통 / 벽 피해 / 탄창', '텅스텐 굴착탄', '파쇄탄 아이콘과 공유하지 않고 관통 화살촉 실루엣을 별도 제작'),
    ('18', '광맥 정제소', '특수 광맥 / 코어 강화', '전투 정제소', '광석 + 정제 장치 조합. 광맥 탐지망과 구분되는 정제/가공 모티프'),
    ('19', '연사 터보', '연사력 / 재장전 보정', '터보 회전자', '회전 약실 아이콘에 속도선 오버레이로 통합 가능; 별도 본체 제작 불필요'),
    ('20', '제어 네트워크', '심층 성장 / 복합 보정', '심층 지배망', '스캐너·전력 노드 계열을 재사용하되, 최종 특성 전용 연결망 오버레이 권장'),
]

inventory = [
    ('일반 · 드릴러', '토크 증폭 모터', '일반 채굴과 기반암 균열 속도 +18%', '01 드릴 모터·비트'),
    ('일반 · 드릴러', '균열 고착 수지', '균열 유지 유예 +1.25초 · 소실 속도 -22%', '02 균열·단층'),
    ('희귀 · 드릴러', '단층 추적 비트', '기반암과 코어 암반 균열 속도 +25%', '01 드릴 모터·비트'),
    ('희귀 · 드릴러', '과압 실린더', 'Q의 균열 압력 +25% · 재사용 대기시간 -15%', '01 드릴 모터·비트'),
    ('영웅 · 드릴러', '지진 공진축', '기반암 돌파 시 주변 일반 벽에 파쇄 충격파', '04 지진 충격파'),
    ('영웅 · 드릴러', '열회수 균열기', '기반암 균열 +20% · 균열 소실 속도 절반', '02 균열·단층'),
    ('전설 · 드릴러', '맨틀 천공 키', '코어 암반 필요 압력 -38% · 모든 기반암 균열 +25%', '03 코어 암반'),
    ('전설 · 드릴러', '영구 단층 각인', '생성한 균열이 더는 소실되지 않으며 균열 속도 +40%', '02 균열·단층'),
    ('일반 · 거너', '고속 약실 순환기', '파쇄탄 재사용 대기시간 -14%', '05 회전 약실·급탄'),
    ('일반 · 거너', '텅스텐 라이너', '파쇄탄의 벽 피해 +20%', '06 파쇄탄·폭파 작약'),
    ('희귀 · 거너', '가변 지연 신관', '자동 기폭 시간 -28% · 부착 즉시 조기 기폭 가능', '07 신관·기폭기'),
    ('희귀 · 거너', '파쇄 확장 슬리브', '파쇄 반경 +1칸 · 외곽 벽 피해 보정', '06 파쇄탄·폭파 작약'),
    ('영웅 · 거너', '원격 과충전 기폭기', 'E 조기 기폭 시 파쇄 위력 +75%', '07 신관·기폭기'),
    ('영웅 · 거너', '대인 파편 재킷', '파쇄탄의 적 피해 +65% · 중화기 피해 +12%', '08 파편·산탄'),
    ('전설 · 거너', '군집 파쇄 작약', '파쇄 반경 3칸 · 벽 피해 +35%', '06 파쇄탄·폭파 작약'),
    ('전설 · 거너', '무정지 발사 사이클', '파쇄탄 재사용 대기시간 대폭 감소 · 벽 피해 +25%', '05 회전 약실·급탄'),
    ('일반 · 스카우트', '고광도 연소제', '플레어 조명 반경과 지속시간 +18%', '09 플레어·광원'),
    ('일반 · 스카우트', '경량 인장 케이블', '그래플 최대 거리 +1칸 · 이동 속도 +5%', '10 그래플·윈치'),
    ('희귀 · 스카우트', '회생 제동 윈치', '그래플 재사용 대기시간 -28%', '10 그래플·윈치'),
    ('희귀 · 스카우트', '다중 스펙트럼 스캐너', '정찰 반경 +1칸 · 신규 구역 탐사 경험치 +1', '11 스캐너·레이더'),
    ('영웅 · 스카우트', '감시자 광학계', '플레어 독립 시야 +2칸 · 정찰 표식 지속 +1.2초', '11 스캐너·레이더'),
    ('영웅 · 스카우트', '원격 측량 릴레이', '플레어 지속시간 +35% · 탐사 경험치 +2', '09 플레어·광원'),
    ('전설 · 스카우트', '수평선 견인 장치', '그래플 최대 거리 9칸 · 재사용 대기시간 -40% · 이동 +12%', '10 그래플·윈치'),
    ('전설 · 스카우트', '인공 태양 플레어', '플레어 반경·지속 +50% · 독립 시야 +3칸 · 정찰 반경 8칸', '09 플레어·광원'),
    ('일반 · 엔지니어', '고밀도 축전지', '전력 노드 지속시간 +25% · 공급 반경 +0.5칸', '12 축전지·전력 노드'),
    ('일반 · 엔지니어', '확장 탄약 호퍼', '새 센트리 탄창 +6 · 설치된 센트리도 즉시 보급', '13 센트리 터렛'),
    ('희귀 · 엔지니어', '능동 냉각 재킷', '센트리 발사 간격 -22%', '13 센트리 터렛'),
    ('희귀 · 엔지니어', '분산 전력 프로토콜', '전력 노드 최대 +1 · 공급 반경 +0.75칸', '12 축전지·전력 노드'),
    ('영웅 · 엔지니어', '쌍중 센트리 설계', '센트리 최대 설치 수 +1 · 지속시간 +25%', '13 센트리 터렛'),
    ('영웅 · 엔지니어', '중량 레일 센트리', '센트리 피해 +45% · 사거리 +1칸 · 탄창 +6 · 연사력 -11%', '13 센트리 터렛'),
    ('전설 · 엔지니어', '폐쇄형 자율 동력로', '센트리가 전력망 밖에서도 가동 · 재장전 시간 -30%', '12 축전지·전력 노드'),
    ('전설 · 엔지니어', '심층 요새 네트워크', '노드·센트리 최대 3 · 공급 반경 6칸 · 센트리 탄창 +12 · 연사 +20%', '12 축전지·전력 노드'),
    ('심층 보상', '쌍동력 융합', '드릴·총기 피해 +35% · 재장전 시간 +12%', '14 쌍동력·동시 운용'),
    ('심층 보상', '파쇄 탄막', '폭발탄과 벽 피해 +60% · 재장전 시간 +18%', '06 파쇄탄·폭파 작약'),
    ('심층 보상', '다중 급탄', '투사체 +1 · 연사력 +15% · 탄창 +3 · 재장전 +10%', '05 회전 약실·급탄'),
    ('심층 보상', '무한 도탄', '도탄 +2 · 재장전 시간 +12%', '15 도탄'),
    ('심층 보상', '굴진 과급', '드릴 피해 +45% · 이동 속도 +12%', '01 드릴 모터·비트'),
    ('심층 보상', '생존 격벽', '최대 체력 +40 · 체력을 추가 회복', '16 생존 격벽'),
    ('트리 · 굴진', '다이아몬드 비트', '드릴 피해 +20%', '01 드릴 모터·비트'),
    ('트리 · 굴진', '단층 파쇄축', '벽 파괴 시 인접 벽 연쇄 피해', '02 균열·단층'),
    ('트리 · 굴진', '지각 붕괴 코어', '벽 10개마다 대형 굴착 폭발 · 드릴 +25%', '04 지진 충격파'),
    ('트리 · 탄도', '텅스텐 굴착탄', '관통 +1 · 벽 피해 +20% · 탄창 +2', '17 관통 굴착탄'),
    ('트리 · 탄도', '병렬 급탄기', '투사체 +1 · 연사력 +10% · 탄창 +3 · 재장전 +8%', '05 회전 약실·급탄'),
    ('트리 · 탄도', '군집 붕괴탄', '모든 탄환이 채굴 폭발 · 총기 +30% · 재장전 +18%', '06 파쇄탄·폭파 작약'),
    ('트리 · 동력', '터보 회전자', '총기 연사력 +18% · 재장전 +8%', '19 연사 터보'),
    ('트리 · 동력', '쌍축 액추에이터', '동시 사용 +35% · 이동 +8% · 재장전 -15%', '14 쌍동력·동시 운용'),
    ('트리 · 동력', '무제한 출력', '드릴·총 +35% · 연사 +15% · 탄창 +6 · 재장전 +20%', '14 쌍동력·동시 운용'),
    ('트리 · 장악', '광맥 탐지망', '경험치 +20% · 추가 코어 확률 +20%', '11 스캐너·레이더'),
    ('트리 · 장악', '전투 정제소', '특수 광맥 체력 +6 · 코어 강화 · 재장전 -12%', '18 광맥 정제소'),
    ('트리 · 장악', '심층 지배망', '경험 +35% · 벽 피해 +50% · 파편탄 8발 · 탄창 +2', '20 제어 네트워크'),
]


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.header_distance = Inches(0.492)
    sec.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles['Normal']
    normal.font.name = 'Calibri'
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in [
        ('Heading 1',16,BLUE,18,10),('Heading 2',13,BLUE,14,7),('Heading 3',12,DARK_BLUE,10,5)
    ]:
        st = styles[name]
        st.font.name = 'Calibri'
        st._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    add_text(title, '땅굴 크루 특성 아이콘 제작 계획', bold=True, color=NAVY, size=24)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    add_text(subtitle, '전체 특성 중복 사용성 검토 및 최종 제작 목록', color=MUTED, size=11)
    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(12)
    add_text(meta, '기준 데이터: tunnel-crew-infinite-mode-v7.1.2.html  |  작성일: 2026-08-26', color=MUTED, size=9)

    add_callout(doc, '결론', '전체 50개 특성을 검토한 결과, 고유 기본 아이콘은 20개로 정리할 수 있습니다. 이미 제작된 4개를 제외하면 추가 제작 대상은 16개이며, 등급·수치·상태 차이는 카드 레이어와 작은 보조 마크로 처리하는 방식을 권장합니다.')

    add_heading(doc, '1. 검토 범위와 제작 원칙', 1)
    add_para(doc, '검토 범위는 최신 게임 데이터의 세 가지 특성 시스템입니다. 일반 레벨업 특성 32개, 심층 보상 특성 6개, 특성 트리 노드 12개를 이름과 효과 기준으로 분류했습니다.')
    make_table(doc, ['구분', '개수', '아이콘 운영 방향'], [
        ('일반 레벨업 특성', '32개', '역할별 장비 아이콘을 재사용하고 등급은 카드 프레임으로 구분'),
        ('심층 보상 특성', '6개', '기존 장비 아이콘 재사용 + 복합 효과 보조 마크'),
        ('특성 트리 노드', '12개', '일반 특성과 같은 장비 아이콘을 재사용하되 트리 전용 오버레이 추가'),
        ('합계', '50개', '고유 기본 아이콘 20개로 커버')
    ], [2600, 1000, 5760], font_size=9)
    add_para(doc, '등급은 아이콘 자체에 굳이 새 그래픽을 추가하지 않는 편이 좋습니다. 일반·희귀·영웅·전설의 온도 차이는 이미 확정한 카드 프레임, 테두리 광택, 배경 링, 파티클 강도로 충분히 전달할 수 있습니다.', after=8)

    add_heading(doc, '2. 최종 제작 대상: 고유 기본 아이콘 20개', 1)
    add_para(doc, '아래 표의 번호를 아이콘 파일명과 구현 데이터의 기준 ID로 사용할 수 있습니다. 한 그룹에 포함된 특성은 같은 기본 아이콘을 사용하고, 괄호 안의 보조 요소만 추가합니다.')
    make_table(doc, ['ID', '기본 아이콘', '대표 의미', '재사용 특성', '제작 메모'], base_groups, [520, 1450, 1500, 3200, 2690], font_size=8)

    add_heading(doc, '3. 이미 제작된 아이콘과 추가 제작량', 1)
    make_table(doc, ['기존 아이콘', '연결 기본 아이콘', '상태'], [
        ('trait-icon-driller-torque-motor-v1.png', '01 드릴 모터·비트', '기존 제작 완료'),
        ('trait-icon-gunner-chamber-cycler-v1.png', '05 회전 약실·급탄', '기존 제작 완료'),
        ('trait-icon-scout-high-luminosity-fuel-v1.png', '09 플레어·광원', '기존 제작 완료'),
        ('trait-icon-engineer-high-density-battery-v1.png', '12 축전지·전력 노드', '기존 제작 완료'),
        ('추가 제작 대상', '02, 03, 04, 06, 07, 08, 10, 11, 13, 14, 15, 16, 17, 18, 19, 20', '16개')
    ], [3300, 3800, 2260], font_size=9)
    add_callout(doc, '제작 우선순위', '먼저 01~13을 제작하면 32개 일반 레벨업 특성 대부분을 커버할 수 있습니다. 이후 14~20을 제작하면 심층 보상과 특성 트리까지 전체 시스템을 완성할 수 있습니다.')

    add_heading(doc, '4. 전체 특성 매핑 목록', 1)
    add_para(doc, '아래 목록은 원본 코드의 이름과 설명을 유지하면서, 최종 제작 기본 아이콘 ID를 연결한 전체 인벤토리입니다. 이 표를 아이콘 제작 체크리스트와 데이터 연결 검수표로 사용할 수 있습니다.')
    make_table(doc, ['시스템 / 역할', '특성명', '효과 설명', '기본 아이콘'], inventory, [1650, 2050, 3750, 1910], font_size=7)

    add_heading(doc, '5. 구현 권장안', 1)
    add_heading(doc, '아이콘 본체와 보조 마크 분리', 2)
    add_para(doc, '아이콘 본체는 장비나 오브젝트의 정체성을 담당하고, 수치나 상태는 작은 보조 마크로 분리합니다. 예를 들어 회전 약실 아이콘에 속도선, 탄환 수, 시계 마크를 겹쳐 쓰면 여러 특성을 한 아이콘으로 처리할 수 있습니다.')
    make_table(doc, ['보조 마크', '표현할 의미', '적용 예'], [
        ('속도선 / 회전 링', '연사력, 재사용 대기시간, 회전 속도', '고속 약실 순환기, 터보 회전자'),
        ('시계 / 신호파', '지연, 원격, 재사용 대기시간', '가변 지연 신관, 원격 과충전 기폭기'),
        ('화살촉 / 관통선', '관통, 거리, 방향성', '텅스텐 굴착탄, 수평선 견인 장치'),
        ('충격파 링', '반경, 광역 피해, 연쇄 파괴', '지진 공진축, 군집 파쇄 작약'),
        ('노드 / 연결선', '설치 수, 공급 반경, 네트워크', '분산 전력 프로토콜, 심층 지배망'),
        ('방패 / 하트', '체력, 회복, 방어', '생존 격벽'),
        ('광맥 점 / 별가루', '경험치, 코어, 정찰 보상', '광맥 탐지망, 원격 측량 릴레이')
    ], [2100, 3300, 3960], font_size=8)

    add_heading(doc, '6. 제작 파일 제안', 1)
    add_para(doc, '아이콘은 카드 중앙 원형 슬롯에 맞춘 1254×1254 정사각형 투명 PNG 규격을 유지합니다. 기본 오브젝트가 캔버스의 약 70%를 차지하도록 안전 여백을 두고, 카드 프레임이나 텍스트를 아이콘 파일에 포함하지 않습니다.')
    make_table(doc, ['구분', '권장 규격 / 규칙'], [
        ('파일 포맷', 'PNG-24 또는 PNG-32, 투명 알파 포함'),
        ('캔버스', '1254×1254 px, 1:1'),
        ('안전 여백', '오브젝트가 캔버스의 약 70% 이내, 외곽 접촉 금지'),
        ('시점', '정면에 가까운 3/4 오브젝트 뷰, 중앙 정렬'),
        ('스타일', '땅굴 크루의 3D 카툰 금속·황동 스타일, 과도한 파티클 금지'),
        ('파일명 예시', 'trait-icon-<base-id>-v1.png 또는 trait-icon-<semantic-name>-v1.png')
    ], [2100, 7260], font_size=9)

    add_para(doc, '원본 기준: tunnel-crew-infinite-mode-v7.1.2.html의 INF_TRAITS, INF_LEGENDS, INF_BRANCHES 정의 및 INF_ICON_KIND 분류 로직.', before=10, after=0, color=MUTED, size=8)

    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_text(footer, 'Tunnel Crew · Trait Icon Production Plan', color=MUTED, size=8)

    doc.save(OUT)


if __name__ == '__main__':
    build()
